"""
Word Operations Helper Scripts

Tool functions for the word_operations skill.
Provides CRUD operations on Word documents.
"""

from pathlib import Path
from typing import Optional, Dict, Any, List


def create_document(file_path: str, title: str = "", author: str = "") -> Dict[str, Any]:
    """
    Create a new Word document with a title paragraph.

    Args:
        file_path: Path where the .docx file will be saved
        title: Document title text
        author: Optional author name

    Returns:
        Dict with status and file info
    """
    from docx import Document
    from docx.shared import Pt

    path = Path(file_path)
    path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    if title:
        heading = doc.add_heading(title, level=0)
    if author:
        doc.core_properties.author = author

    doc.save(str(path))

    return {
        "status": "created",
        "file_path": str(path.resolve()),
        "title": title,
        "size_bytes": path.stat().st_size,
    }


def add_paragraph(file_path: str, text: str, style: str = "Normal", bold: bool = False, font_size: int = None) -> Dict[str, Any]:
    """
    Add a paragraph to the end of a Word document.

    Args:
        file_path: Path to the .docx file
        text: Paragraph text
        style: Paragraph style (Normal, Heading 1, Heading 2, List Bullet, etc.)
        bold: Make text bold
        font_size: Font size in points

    Returns:
        Dict with status
    """
    from docx import Document
    from docx.shared import Pt

    path = Path(file_path)
    if not path.exists():
        return {"error": f"File not found: {file_path}"}

    doc = Document(str(path))
    para = doc.add_paragraph(text, style=style)

    if bold or font_size:
        for run in para.runs:
            if bold:
                run.bold = True
            if font_size:
                run.font.size = Pt(font_size)

    doc.save(str(path))

    return {"status": "added", "text": text[:80], "style": style}


def add_heading(file_path: str, text: str, level: int = 1) -> Dict[str, Any]:
    """
    Add a heading to the document.

    Args:
        file_path: Path to the .docx file
        text: Heading text
        level: Heading level (1-9, where 1 is largest)

    Returns:
        Dict with status
    """
    from docx import Document

    path = Path(file_path)
    if not path.exists():
        return {"error": f"File not found: {file_path}"}

    doc = Document(str(path))
    doc.add_heading(text, level=level)
    doc.save(str(path))

    return {"status": "added", "heading": text, "level": level}


def add_table(file_path: str, headers: List[str], rows: List[List[str]], style: str = "Table Grid") -> Dict[str, Any]:
    """
    Add a table to the document.

    Args:
        file_path: Path to the .docx file
        headers: List of header column names
        rows: 2D list of row data
        style: Table style name (Table Grid, Plain Text, Light List, etc.)

    Returns:
        Dict with status and table info
    """
    from docx import Document

    path = Path(file_path)
    if not path.exists():
        return {"error": f"File not found: {file_path}"}

    doc = Document(str(path))
    table = doc.add_table(rows=1 + len(rows), cols=len(headers), style=style)

    # Headers
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            table.rows[r_idx + 1].cells[c_idx].text = str(value)

    doc.save(str(path))

    return {"status": "added", "rows": len(rows), "columns": len(headers)}


def read_document(file_path: str, max_chars: int = 5000) -> Dict[str, Any]:
    """
    Read text content from a Word document.

    Args:
        file_path: Path to the .docx file
        max_chars: Maximum characters to return

    Returns:
        Dict with text content and metadata
    """
    from docx import Document

    path = Path(file_path)
    if not path.exists():
        return {"error": f"File not found: {file_path}"}

    doc = Document(str(path))
    paragraphs = []
    total_chars = 0

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            if total_chars + len(text) > max_chars:
                remaining = max_chars - total_chars
                paragraphs.append(text[:remaining] + "...")
                break
            paragraphs.append(text)
            total_chars += len(text)

    tables = []
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            table_data.append(row_data)
        tables.append(table_data)

    word_count = sum(len(p.split()) for p in paragraphs)

    return {
        "file_path": str(path.resolve()),
        "paragraph_count": len(doc.paragraphs),
        "table_count": len(doc.tables),
        "word_count": word_count,
        "text": "\n\n".join(paragraphs),
        "tables": tables,
        "truncated": total_chars >= max_chars,
    }


def get_document_info(file_path: str) -> Dict[str, Any]:
    """
    Get metadata about a Word document.

    Args:
        file_path: Path to the .docx file

    Returns:
        Dict with metadata
    """
    from docx import Document

    path = Path(file_path)
    if not path.exists():
        return {"error": f"File not found: {file_path}"}

    doc = Document(str(path))
    word_count = sum(len(p.text.split()) for p in doc.paragraphs)

    return {
        "file_path": str(path.resolve()),
        "title": doc.core_properties.title or "",
        "author": doc.core_properties.author or "",
        "paragraph_count": len(doc.paragraphs),
        "table_count": len(doc.tables),
        "word_count": word_count,
        "size_bytes": path.stat().st_size,
    }


def replace_text(file_path: str, old_text: str, new_text: str) -> Dict[str, Any]:
    """
    Replace all occurrences of text in the document.

    Args:
        file_path: Path to the .docx file
        old_text: Text to find
        new_text: Text to replace with

    Returns:
        Dict with count of replacements
    """
    from docx import Document

    path = Path(file_path)
    if not path.exists():
        return {"error": f"File not found: {file_path}"}

    doc = Document(str(path))
    count = 0

    for para in doc.paragraphs:
        if old_text in para.text:
            for run in para.runs:
                if old_text in run.text:
                    run.text = run.text.replace(old_text, new_text)
                    count += 1

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if old_text in cell.text:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if old_text in run.text:
                                run.text = run.text.replace(old_text, new_text)
                                count += 1

    doc.save(str(path))

    return {"status": "replaced", "count": count, "old_text": old_text, "new_text": new_text}


def set_margins(file_path: str, top: float = 1.0, bottom: float = 1.0, left: float = 1.0, right: float = 1.0) -> Dict[str, Any]:
    """
    Set page margins for the document.

    Args:
        file_path: Path to the .docx file
        top: Top margin in inches
        bottom: Bottom margin in inches
        left: Left margin in inches
        right: Right margin in inches

    Returns:
        Dict with status
    """
    from docx import Document
    from docx.shared import Inches

    path = Path(file_path)
    if not path.exists():
        return {"error": f"File not found: {file_path}"}

    doc = Document(str(path))
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)

    doc.save(str(path))

    return {"status": "margins_set", "top": top, "bottom": bottom, "left": left, "right": right}


def add_image(file_path: str, image_path: str, width: float = 4.0) -> Dict[str, Any]:
    """
    Add an image to the end of the document.

    Args:
        file_path: Path to the .docx file
        image_path: Path to the image file
        width: Width in inches (height scales proportionally)

    Returns:
        Dict with status
    """
    from docx import Document
    from docx.shared import Inches

    path = Path(file_path)
    if not path.exists():
        return {"error": f"File not found: {file_path}"}

    img_path = Path(image_path)
    if not img_path.exists():
        return {"error": f"Image not found: {image_path}"}

    doc = Document(str(path))
    doc.add_picture(str(img_path), width=Inches(width))
    doc.save(str(path))

    return {"status": "image_added", "image": str(img_path.name), "width_inches": width}
