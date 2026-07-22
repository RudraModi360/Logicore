"""
DOCX auto-fixes.

Handles fixing common Word document issues:
- Add headings to headingless documents
- Remove excessive empty paragraphs
- Fix table formatting
"""

from __future__ import annotations

import logging
import os
from typing import TYPE_CHECKING

if TYPE_CHECKING:
    from logicore.verification.auto_fix import AutoFixEngine

logger = logging.getLogger(__name__)


def register_fixes(engine: "AutoFixEngine") -> None:
    """Register DOCX fix handlers with the auto-fix engine."""
    engine.register_fix("structure", "no headings", _fix_add_headings)
    engine.register_fix("content", "empty paragraphs", _fix_empty_paragraphs)
    engine.register_fix("content", "all empty cells", _fix_empty_table)


def _fix_add_headings(artifact_path: str, issue) -> bool:
    """Add heading style to the first paragraph if document has no headings."""
    try:
        from docx import Document
    except ImportError:
        return False

    try:
        doc = Document(artifact_path)

        # Find first non-empty paragraph.
        for para in doc.paragraphs:
            if para.text.strip():
                # Apply Heading 1 style.
                para.style = doc.styles["Heading 1"]
                doc.save(artifact_path)
                return True

        return False

    except Exception as exc:
        logger.debug(f"Failed to add headings: {exc}")
        return False


def _fix_empty_paragraphs(artifact_path: str, issue) -> bool:
    """Remove excessive empty paragraphs (keep max 2 consecutive)."""
    try:
        from docx import Document
    except ImportError:
        return False

    try:
        doc = Document(artifact_path)

        # Find and remove consecutive empty paragraphs.
        paragraphs_to_remove = []
        consecutive_empty = 0

        for i, para in enumerate(doc.paragraphs):
            if not para.text.strip() and not para.runs:
                consecutive_empty += 1
                if consecutive_empty > 2:
                    paragraphs_to_remove.append(para)
            else:
                consecutive_empty = 0

        if not paragraphs_to_remove:
            return False

        # Remove paragraphs (from end to start to preserve indices).
        for para in reversed(paragraphs_to_remove):
            p_element = para._element
            p_element.getparent().remove(p_element)

        doc.save(artifact_path)
        return True

    except Exception as exc:
        logger.debug(f"Failed to remove empty paragraphs: {exc}")
        return False


def _fix_empty_table(artifact_path: str, issue) -> bool:
    """Remove tables with all empty cells."""
    try:
        from docx import Document
    except ImportError:
        return False

    try:
        doc = Document(artifact_path)

        tables_to_remove = []

        for table in doc.tables:
            # Check if all cells are empty.
            all_empty = True
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        all_empty = False
                        break
                if not all_empty:
                    break

            if all_empty:
                tables_to_remove.append(table)

        if not tables_to_remove:
            return False

        # Remove tables.
        for table in tables_to_remove:
            t_element = table._tbl
            t_element.getparent().remove(t_element)

        doc.save(artifact_path)
        return True

    except Exception as exc:
        logger.debug(f"Failed to remove empty tables: {exc}")
        return False
