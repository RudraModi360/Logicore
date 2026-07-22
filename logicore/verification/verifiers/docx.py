"""
DOCX verifier.

Verifies Word documents: structure, content, formatting, tables.
"""

from __future__ import annotations

import os
from typing import List, Optional, Set

from logicore.verification.base_verifier import BaseVerifier
from logicore.verification.result import VerificationIssue


# Supported DOCX extensions.
DOCX_EXTENSIONS: Set[str] = {".docx", ".doc"}

# Maximum paragraphs for a reasonable document.
MAX_PARAGRAPHS = 5000

# Maximum tables.
MAX_TABLES = 100


class DOCXVerifier(BaseVerifier):
    """Verify Word documents: structure, content, formatting, tables.

    Checks:
    - Document can be opened and parsed
    - Has meaningful content (paragraphs, tables, or images)
    - Has headings for structure
    - Tables are properly formatted
    - Images are present if expected
    """

    def supported_extensions(self) -> Set[str]:
        return DOCX_EXTENSIONS

    def _verify_content(
        self,
        artifact_path: str,
        issues: List[VerificationIssue],
        requirements: Optional[str],
    ) -> None:
        ext = os.path.splitext(artifact_path)[1].lower()

        if ext == ".doc":
            # .doc files need different handling (not ZIP-based).
            self._verify_doc(artifact_path, issues)
        else:
            self._verify_docx(artifact_path, issues, requirements)

    def _verify_docx(
        self,
        artifact_path: str,
        issues: List[VerificationIssue],
        requirements: Optional[str],
    ) -> None:
        """Verify DOCX files using python-docx."""
        try:
            from docx import Document
        except ImportError:
            self._add_issue(
                issues,
                severity="info",
                category="dependency",
                description="python-docx not installed — skipping detailed DOCX verification",
            )
            return

        try:
            doc = Document(artifact_path)
        except Exception as exc:
            self._add_issue(
                issues,
                severity="critical",
                category="corruption",
                description=f"Cannot open DOCX file: {exc}",
                auto_fixable=False,
            )
            return

        # Content check: has paragraphs or tables.
        has_paragraphs = len(doc.paragraphs) > 0
        has_tables = len(doc.tables) > 0
        has_images = len(doc.inline_shapes) > 0

        if not has_paragraphs and not has_tables and not has_images:
            self._add_issue(
                issues,
                severity="critical",
                category="content",
                description="Document has no content (no paragraphs, tables, or images)",
            )
            return

        # Paragraph count warning.
        if len(doc.paragraphs) > MAX_PARAGRAPHS:
            self._add_issue(
                issues,
                severity="warning",
                category="structure",
                description=f"Document has {len(doc.paragraphs)} paragraphs (unusually large)",
                auto_fixable=False,
            )

        # Heading structure check.
        has_heading = False
        for para in doc.paragraphs:
            if para.style and para.style.name.startswith("Heading"):
                has_heading = True
                break

        if not has_paragraphs:
            pass  # No paragraphs to check.
        elif not has_heading and len(doc.paragraphs) > 5:
            self._add_issue(
                issues,
                severity="warning",
                category="structure",
                description="Document has no headings (poor structure)",
                auto_fixable=True,
                fix_suggestion="Add headings to improve document navigation",
            )

        # Empty paragraph check.
        empty_count = 0
        for para in doc.paragraphs:
            if not para.text.strip() and not para.runs:
                empty_count += 1

        if empty_count > len(doc.paragraphs) * 0.5 and len(doc.paragraphs) > 3:
            self._add_issue(
                issues,
                severity="warning",
                category="content",
                description=f"Document has {empty_count} empty paragraphs ({empty_count * 100 // len(doc.paragraphs)}%)",
                auto_fixable=True,
                fix_suggestion="Remove excessive empty paragraphs",
            )

        # Table checks.
        if has_tables:
            self._check_tables(doc, issues)

        # Image check based on requirements.
        if requirements and "image" in requirements.lower() and not has_images:
            self._add_issue(
                issues,
                severity="warning",
                category="content",
                description="No images found in document (user requested images)",
            )

    def _check_tables(
        self,
        doc,
        issues: List[VerificationIssue],
    ) -> None:
        """Check table formatting and content."""
        for i, table in enumerate(doc.tables):
            # Empty table check.
            if len(table.rows) == 0:
                self._add_issue(
                    issues,
                    severity="warning",
                    category="content",
                    description=f"Table {i + 1} is empty",
                    location=f"table {i + 1}",
                )
                continue

            # Single row/column check.
            if len(table.rows) == 1 and len(table.columns) == 1:
                self._add_issue(
                    issues,
                    severity="info",
                    category="content",
                    description=f"Table {i + 1} has only one cell",
                    location=f"table {i + 1}",
                )

            # Empty cells check.
            total_cells = 0
            empty_cells = 0
            for row in table.rows:
                for cell in row.cells:
                    total_cells += 1
                    if not cell.text.strip():
                        empty_cells += 1

            if total_cells > 0 and empty_cells == total_cells:
                self._add_issue(
                    issues,
                    severity="warning",
                    category="content",
                    description=f"Table {i + 1} has all empty cells",
                    location=f"table {i + 1}",
                )

    def _verify_doc(
        self,
        artifact_path: str,
        issues: List[VerificationIssue],
    ) -> None:
        """Basic verification for legacy .doc files."""
        # .doc files are binary and harder to parse without specialized libraries.
        # We can only do basic checks.
        try:
            with open(artifact_path, "rb") as f:
                header = f.read(8)
        except Exception as exc:
            self._add_issue(
                issues,
                severity="critical",
                category="corruption",
                description=f"Cannot read .doc file: {exc}",
            )
            return

        # OLE compound document magic bytes.
        if not header.startswith(b"\xd0\xcf\x11\xe0"):
            self._add_issue(
                issues,
                severity="critical",
                category="format",
                description="File does not have a valid .doc header",
            )


def get_verifier() -> DOCXVerifier:
    """Return a DOCXVerifier instance for registry auto-discovery."""
    return DOCXVerifier()
